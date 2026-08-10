"""
Sinh HƯỚNG DẪN SỬ DỤNG KHO (Word .docx) — dành riêng cho người làm kho.

Khác gen_huong_dan.py (bao cả CRM/CEO, viết cho người đã quen hệ thống): file
này chỉ nói về KHO, viết theo lối "đóng vai" — Nhân viên kho và Kho trưởng mỗi
ngày làm gì, bấm ở đâu, ra kết quả gì, ai làm tiếp.

Chạy:  python scripts/gen_huong_dan_kho.py [output.docx]
"""
from __future__ import annotations

import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

FLAME = RGBColor(0xE2, 0x5A, 0x1C)
BLUE = RGBColor(0x12, 0x4D, 0xB5)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
RED = RGBColor(0xB0, 0x2A, 0x1F)
GREY = RGBColor(0x33, 0x33, 0x33)


def build(out: str):
    d = Document()
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(10.5)

    def h(text, level=1):
        p = d.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = FLAME if level <= 1 else GREY
        return p

    def para(text=''):
        d.add_paragraph(text)

    def kv(label, text):
        p = d.add_paragraph()
        r = p.add_run(label + ': ')
        r.bold = True
        r.font.color.rgb = FLAME
        p.add_run(text)

    def flow(text):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = BLUE

    def note(text, color=GREEN, icon='✔'):
        p = d.add_paragraph()
        r = p.add_run(f'{icon} {text}')
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = color

    def warn(text):
        note(text, RED, '⚠')

    def steps(items):
        for it in items:
            d.add_paragraph(it, style='List Number')

    def bullets(items):
        for it in items:
            d.add_paragraph(it, style='List Bullet')

    def table(headers, rows, widths=None):
        t = d.add_table(rows=1, cols=len(headers))
        t.style = 'Light Grid Accent 1'
        for i, hd in enumerate(headers):
            t.rows[0].cells[i].text = hd
            for r in t.rows[0].cells[i].paragraphs[0].runs:
                r.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9.5)
        d.add_paragraph()
        return t

    def buttons(rows):
        table(['Nút / Ô nhập', 'Bấm vào để làm gì'], rows)

    def scenario(title, rows):
        p = d.add_paragraph()
        r = p.add_run('▸ ' + title)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = GREEN
        table(['#', 'Ai làm', 'Làm gì', 'Xong thì thành ra sao', 'Ai làm tiếp'], rows)

    # ════════════════════ BÌA ════════════════════
    t = d.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('HƯỚNG DẪN SỬ DỤNG KHO\nTOKINARC WMS')
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = FLAME
    s = d.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run('Dành cho Nhân viên kho và Kho trưởng — từ đầu tới cuối').italic = True
    para()
    para('Tài liệu này viết theo lối kể việc: mỗi ngày anh/chị làm gì, bấm ở đâu, '
         'bấm xong ra sao, và ai làm tiếp. Không cần biết gì về máy tính ngoài việc '
         'dùng được trình duyệt và điện thoại có camera.')
    para()
    kv('Cách đọc nhanh', 'Xem Phần 2 (việc hằng ngày) là đủ chạy được. '
       'Phần 3 tra cứu từng trang khi cần. Phần 4 xử lý khi gặp trục trặc.')
    d.add_page_break()

    # ════════════════════ 1. BẮT ĐẦU ════════════════════
    h('Phần 1. Bắt đầu — hiểu hệ thống trong 5 phút', 1)

    h('1.1. Hai vai trò trong kho', 2)
    para('Hệ thống chia người làm kho thành 2 vai. Đăng nhập bằng tài khoản nào thì '
         'thấy đúng phần việc của vai đó.')
    table(['Vai trò', 'Là ai', 'Được làm', 'KHÔNG được làm'], [
        ['Nhân viên kho\n(NV Kho)', 'Người trực tiếp bốc xếp, nhận hàng, soạn hàng',
         'Nhập kho, xuất kho, soạn hàng, quét mã, đếm kiểm kê, gán mã QR, '
         'thêm/sửa sản phẩm trong danh mục',
         'Sửa số tồn kho bằng tay, duyệt kiểm kê, tạo/sửa kho–khu–ô'],
        ['Kho trưởng\n(Quản lý kho)', 'Người chịu trách nhiệm số liệu kho',
         'Tất cả việc của NV Kho, cộng thêm: điều chỉnh tồn, duyệt kiểm kê, '
         'tạo/sửa kho–khu–ô, quản lý nhóm hàng, đơn mua & nhà cung cấp',
         '—'],
    ])
    warn('Ranh giới quan trọng: NV Kho KHÔNG tự sửa được số tồn. Muốn đổi số tồn '
         'phải qua kiểm kê rồi Kho trưởng duyệt. Đây là cố ý — để số liệu kho luôn '
         'có người chịu trách nhiệm, không ai sửa lén được.')

    h('1.2. Kho được chia thế nào', 2)
    para('Hàng không để chung một đống mà chia theo 4 tầng, giống như địa chỉ nhà:')
    flow('KHO ─► KHU (Zone) ─► KỆ (Rack) ─► Ô (Bin)')
    table(['Tầng', 'Nghĩa', 'Ví dụ'], [
        ['Kho', 'Cả nhà kho', 'Kho 1, Kho HCM'],
        ['Khu', 'Một khu vực trong kho', 'Khu A (ngoài), Khu B (súng hàn)'],
        ['Kệ', 'Một giá kệ, có thể nhiều tầng', 'Kệ 1-1, Kệ 1-1 Tầng 2'],
        ['Ô', 'Chỗ đặt hàng cụ thể nhất', '1-A-1-1-01'],
    ])
    kv('Mã ô đầy đủ', 'Ghép cả 4 tầng lại, ví dụ 1-A-1-1-01 nghĩa là Kho 1 → Khu A → '
       'Kệ 1-1 → Ô 01. Nhìn mã là biết đi tới đâu lấy hàng.')
    note('Một ô chứa được NHIỀU mặt hàng khác nhau — không bắt buộc mỗi ô một loại.')

    h('1.3. Ba con số hay gặp', 2)
    table(['Cột', 'Nghĩa'], [
        ['Tồn', 'Số lượng đang thực có trong ô đó'],
        ['Giữ', 'Số đã bị một đơn xuất "xí phần", chưa lấy đi nhưng không được dùng cho đơn khác'],
        ['Còn dùng được', 'Tồn trừ Giữ — số thật sự còn rảnh để bán/xuất'],
    ])

    h('1.4. Đăng nhập', 2)
    steps([
        'Mở trình duyệt, vào địa chỉ hệ thống công ty cấp.',
        'Nhập tài khoản và mật khẩu → bấm "Đăng nhập".',
        'Vào được rồi thì menu bên trái là toàn bộ việc của kho.',
    ])
    note('Điện thoại dùng được bình thường — nên dùng điện thoại khi cần quét mã, '
         'vì có camera sẵn trong tay.')
    d.add_page_break()

    # ════════════════════ 2. VIỆC HẰNG NGÀY ════════════════════
    h('Phần 2. Việc hằng ngày — làm theo là chạy', 1)
    para('Phần này là xương sống. Mỗi mục là một tình huống có thật trong kho, '
         'kể từ lúc bắt đầu tới lúc xong.')

    # ---- 2.1 NHẬP KHO ----
    h('2.1. Hàng về kho (Nhập kho)', 2)
    flow('Hàng tới ─► Tạo đơn nhập ─► Quét/nhập từng mặt hàng ─► Chọn ô cất ─►')
    flow('   Xác nhận nhận hàng ─► Tồn kho tự cộng lên')
    kv('Ai làm', 'Nhân viên kho')
    kv('Vào đâu', 'Menu trái → Nhập kho → nút "Tạo đơn nhập"')

    para('Hai loại đơn nhập, chọn đúng tab trước khi tạo:')
    table(['Tab', 'Dùng khi nào'], [
        ['Nội bộ', 'Hàng chuyển từ kho khác về, hàng trả lại, hàng nội bộ'],
        ['Nhà cung cấp (NCC)', 'Hàng mua từ nhà cung cấp, có hoá đơn'],
    ])

    para('Các bước:')
    steps([
        'Bấm "Tạo đơn nhập".',
        'Chọn kho nhận. Nếu là hàng nhà cung cấp thì chọn thêm nhà cung cấp, '
        'số hoá đơn, ngày hoá đơn.',
        'Bấm "Quét mã" để bật camera, rồi đưa tem trên thùng hàng vào khung. '
        'Quét trúng là hệ thống tự thêm một dòng hàng.',
        'Nếu có 5 thùng cùng loại thì quét 5 lần — mỗi lần quét cộng thêm 1 vào số lượng.',
        'Không quét được thì gõ tay: bấm ô "Mặt hàng" rồi gõ mã hoặc tên để tìm.',
        'Điền "SL" (số lượng thực nhận) và "Đơn giá nhập" nếu có hoá đơn.',
        'Chọn "Bin đích" — tức là sẽ cất vào ô nào. Bấm vào ô này rồi gõ để tìm.',
        'Xem lại khối "Xem trước" ở dưới cho khớp thực tế rồi bấm "Tạo".',
        'Về danh sách, tìm đơn vừa tạo, bấm nút dấu tích xanh "Nhận đủ" '
        '— lúc này tồn kho mới thật sự cộng lên.',
    ])
    warn('Nhớ chọn Bin đích TRƯỚC khi bấm Tạo. Không chọn thì hàng vào kho nhưng '
         'không biết nằm ở đâu, sau này đi tìm rất mất công.')
    warn('Số lượng tự nhảy khi quét: camera quét liên tục, mã rời khỏi khung rồi '
         'quay lại là tính thêm một lần quét. Rung tay hay loá đèn cũng có thể bị '
         'cộng dư. LUÔN nhìn lại số lượng trước khi bấm Tạo.')

    para('Nếu hàng về thiếu:')
    steps([
        'Bấm nút "Nhận một phần" (thay vì "Nhận đủ").',
        'Ghi lý do thiếu vào ô ghi chú, ví dụ "Còn thiếu 20 cái, NCC hẹn tuần sau".',
        'Phiếu vẫn mở, khi hàng về nốt thì vào nhận tiếp.',
    ])

    buttons([
        ['Tạo đơn nhập', 'Mở form tạo phiếu nhập mới'],
        ['Quét mã', 'Bật camera quét tem — quét ra là tự thêm dòng hàng'],
        ['Tải ảnh lên', 'Chụp sẵn rồi tải ảnh tem lên, dùng khi camera khó bắt'],
        ['Xem (hình con mắt)', 'Xem chi tiết các dòng hàng trong phiếu'],
        ['Nhận đủ (dấu tích xanh)', 'Xác nhận đã nhận đủ hàng → tồn kho cộng lên'],
        ['Nhận một phần', 'Nhận thiếu, phiếu vẫn mở để nhận tiếp'],
        ['Sửa (bút chì)', 'Sửa phiếu, chỉ sửa được khi phiếu còn ở trạng thái Nháp'],
        ['Sửa ngày nhập kho (lịch)', 'Chỉ đổi ngày, không đụng số lượng — dùng khi ghi nhầm ngày'],
        ['Xoá (thùng rác)', 'Ẩn phiếu đi, không mất dữ liệu'],
        ['Khôi phục', 'Lấy lại phiếu đã lỡ xoá'],
        ['Bin đích', 'Cho biết đơn này cất vào ô nào; nhiều ô thì hiện "2 vị trí"'],
    ])

    scenario('Ví dụ: nhận 5 thùng béc hàn từ nhà cung cấp', [
        ['1', 'NV Kho', 'Tab "Nhà cung cấp" → Tạo đơn nhập → chọn NCC, số hoá đơn',
         'Phiếu trắng đang mở', 'NV Kho'],
        ['2', 'NV Kho', 'Bấm Quét mã, quét tem 5 thùng',
         'Dòng hàng hiện ra, SL = 5', 'NV Kho'],
        ['3', 'NV Kho', 'Chọn Bin đích 1-A-1-1-01, điền đơn giá',
         'Xem trước hiện đủ', 'NV Kho'],
        ['4', 'NV Kho', 'Bấm "Tạo"', 'Phiếu trạng thái Nháp', 'NV Kho'],
        ['5', 'NV Kho', 'Xếp hàng lên kệ đúng ô, rồi bấm "Nhận đủ"',
         'Tồn kho ô đó +5, phiếu chuyển Đã cất kho', 'Xong'],
    ])

    # ---- 2.2 XUẤT KHO ----
    h('2.2. Hàng đi khỏi kho (Xuất kho)', 2)
    flow('Có đơn cần giao ─► Tạo/nhận phiếu xuất ─► Soạn hàng (quét) ─► Giao hàng')
    kv('Ai làm', 'Nhân viên kho')
    kv('Vào đâu', 'Menu trái → Xuất kho')

    para('Phiếu xuất tới từ 2 đường: bên bán hàng bấm "Giao" thì phiếu tự sinh ra, '
         'hoặc kho tự bấm "Tạo đơn xuất" khi xuất nội bộ.')
    steps([
        'Mở phiếu xuất cần làm.',
        'Bấm "Quét" để soạn hàng. Hệ thống đã tự tính sẵn lấy hàng từ ô nào.',
        'Đi tới đúng ô đó, quét mã hàng để xác nhận đã lấy.',
        'Lấy đủ hết các dòng thì bấm "Giao".',
        'Xong: tồn kho trừ đi, phiếu chuyển sang Đã giao.',
    ])
    note('Giao thiếu, hàng còn lại giao sau: nút sẽ đổi thành "Giao tiếp" — bấm '
         'tiếp khi giao nốt phần còn lại.')
    kv('Hệ thống chọn ô lấy hàng thế nào', 'Tự chọn theo quy tắc của đơn: '
       'FIFO (hàng vào trước lấy trước) · FEFO (hàng hết hạn trước lấy trước) · '
       'NEAREST (ô gần cửa xuất nhất, đi ít bước nhất).')
    note('Trong lúc soạn, số hàng đó bị chuyển sang cột "Giữ" — đơn khác không '
         'lấy trùng được nữa.')
    buttons([
        ['Tạo đơn xuất', 'Tạo phiếu xuất mới (xuất nội bộ, không qua bán hàng)'],
        ['Xem (con mắt)', 'Xem các dòng hàng trong phiếu'],
        ['Quét', 'Soạn hàng — quét xác nhận đã lấy từng món'],
        ['Giao', 'Chốt đã giao → trừ tồn kho (giao thiếu thì nút thành "Giao tiếp")'],
        ['Sửa / Xoá', 'Chỉnh hoặc ẩn phiếu khi chưa giao'],
    ])

    # ---- 2.3 TỒN KHO ----
    h('2.3. Xem còn bao nhiêu hàng (Tồn kho)', 2)
    kv('Ai làm', 'Cả hai vai đều xem được. Chỉ Kho trưởng mới sửa được số.')
    kv('Vào đâu', 'Menu trái → Tồn kho')
    para('Mỗi dòng là một mặt hàng nằm ở một ô cụ thể. Cùng một mặt hàng để ở 2 ô '
         'thì hiện thành 2 dòng — nhìn cột "Vị trí" để phân biệt.')
    buttons([
        ['Ô tìm kiếm', 'Gõ mã hoặc tên hàng để lọc'],
        ['Xem theo nhóm hàng', 'Gộp lại theo nhóm, xem tổng quan thay vì từng dòng'],
        ['Chỉ sắp hết', 'Lọc riêng những mặt hàng dưới mức tối thiểu; số trong ngoặc là bao nhiêu mã đang thiếu'],
        ['Điều chỉnh tồn (thanh trượt)', 'Sửa số tồn — CHỈ Kho trưởng'],
        ['Chuyển kho (hai mũi tên)', 'Chuyển hàng từ ô này sang ô khác'],
        ['Xuất Excel theo nhóm', 'Tải báo cáo tồn kho theo nhóm hàng'],
    ])
    warn('NV Kho bấm "Điều chỉnh tồn" sẽ bị báo "Chỉ Quản lý kho trở lên được '
         'điều chỉnh tồn". Đây không phải lỗi — muốn sửa số thì làm kiểm kê.')

    # ---- 2.4 KIỂM KÊ ----
    h('2.4. Đếm hàng thực tế (Kiểm kê)', 2)
    flow('Kho trưởng mở phiên ─► NV Kho đi đếm & quét ─► Kho trưởng duyệt ─► Tồn kho khớp lại')
    kv('Vào đâu', 'Menu trái → Kiểm kê')
    para('Đây là cách DUY NHẤT hợp lệ để sửa số tồn khi đếm thực tế lệch với hệ thống.')
    steps([
        '(NV Kho hoặc Kho trưởng) Bấm "Phiên mới", chọn kho cần đếm.',
        '(NV Kho) Mở phiên đang mở, đi tới từng ô.',
        '(NV Kho) Quét mã hàng, gõ số đếm được thực tế vào, bấm thêm dòng.',
        '(NV Kho) Đếm xong hết thì báo Kho trưởng.',
        '(Kho trưởng) Xem lại các dòng lệch, thấy hợp lý thì bấm "Áp dụng".',
        'Hệ thống sửa tồn kho theo số đã đếm, ghi lại ai duyệt, lúc nào.',
    ])
    note('NV Kho mở phiên sẽ thấy dòng chữ "Chờ Quản lý kho duyệt" thay cho nút '
         'Áp dụng — đếm cứ đếm, duyệt là việc của Kho trưởng.')

    # ---- 2.5 GÁN MÃ ----
    h('2.5. Dán mã cho hàng mới (Gán mã vạch/QR)', 2)
    kv('Ai làm', 'Nhân viên kho')
    kv('Vào đâu', 'Menu trái → Gán mã vạch/QR')
    para('Mục đích: để lần sau quét tem là hệ thống nhận ra ngay đó là hàng gì. '
         'Mỗi sản phẩm gán được tối đa 2 mã — 1 mã QR và 1 mã vạch (barcode).')

    para('Cách 1 — quét từng tem một:')
    steps([
        'Bấm "Bật camera quét" (hoặc "Tải ảnh lên" nếu đã chụp sẵn).',
        'Đưa tem vào khung. Hệ thống đọc xong sẽ báo một trong các trường hợp bên dưới.',
    ])
    table(['Hệ thống báo gì', 'Nghĩa là', 'Anh/chị làm gì'], [
        ['Hiện luôn thông tin sản phẩm', 'Mã này đã gán rồi', 'Không cần làm gì'],
        ['"Sản phẩm X đã có sẵn trong hệ thống"', 'Hàng có trong danh mục, chỉ chưa gán mã',
         'Bấm "Gán mã QR này cho sản phẩm"'],
        ['"Ảnh này có 2 mã…"', 'Tem có cả QR lẫn mã vạch, 1 cái đã gán',
         'Bấm "Gán luôn" để gán nốt cái còn lại'],
        ['"Không tìm thấy… có thể chưa gán"', 'Hàng chưa có trong danh mục',
         'Bấm "Gán mã này cho sản phẩm" rồi chọn hàng, hoặc "Thêm mới" nếu thật sự chưa có'],
    ])

    para('Cách 2 — gán hàng loạt nhiều tem một lúc:')
    steps([
        'Bấm "Gán hàng loạt nhiều ảnh" ở góc trên bên phải.',
        'Chọn tối đa 10 ảnh tem đã chụp sẵn.',
        'Đợi đọc xong, hệ thống chia sẵn thành các nhóm: đã khớp / sẽ tự gán / '
        'cần tự chọn / ảnh mờ không đọc được / mã trùng.',
        'Xử lý nhóm "cần tự chọn": bấm nút gợi ý, hoặc tự chọn sản phẩm.',
        'Kiểm lại rồi bấm "Lưu tất cả" — TRƯỚC khi bấm nút này thì chưa có gì được lưu.',
    ])
    note('Tem có cả QR lẫn mã vạch: chọn sản phẩm cho một mã thì mã còn lại trên '
         'cùng tấm ảnh tự điền theo, khỏi phải chọn hai lần.')
    warn('Tem súng hàn quét chậm hơn tem béc hàn vì mã QR chứa nhiều thông tin gấp '
         '8 lần nên ô mã nhỏ hơn. Đưa camera gần lại khoảng một gang tay và giữ yên; '
         'khó quá thì chụp ảnh rồi dùng "Tải ảnh lên".')

    # ---- 2.6 TÌM HÀNG ----
    h('2.6. Đi tìm hàng nằm ở đâu (Bản đồ kho)', 2)
    kv('Vào đâu', 'Menu trái → Bản đồ kho')
    para('Trang này vẽ lại kho theo đúng hình dạng thật: từng khu, từng kệ, từng '
         'tầng, từng ô. Ô màu cam là có hàng, ô xám là trống.')
    steps([
        'Gõ mã hoặc tên hàng vào ô tìm kiếm → ô chứa hàng đó sáng viền vàng, '
        'màn hình tự cuộn tới.',
        'Bấm vào một Ô → hiện danh sách mọi mặt hàng trong ô đó.',
        'Bấm vào chữ "Tầng" → hiện toàn bộ hàng trên tầng, kèm cột cho biết nằm ô nào.',
        'Bấm vào chữ "Kệ" → hiện toàn bộ hàng cả kệ.',
        'Muốn in ra giấy hay gửi cho ai thì bấm "Xuất Excel" ngay trong cửa sổ đó.',
    ])
    note('Ô nào ghi "3 mã" nghĩa là ô đó đang chứa 3 mặt hàng khác nhau — bấm vào '
         'xem đủ cả ba.')

    # ---- 2.7 TRUY XUẤT + LỊCH SỬ ----
    h('2.7. Tra ngược một món hàng (Truy xuất & Lịch sử kho)', 2)
    table(['Trang', 'Dùng khi nào', 'Cho biết gì'], [
        ['Truy xuất (Serial/Lô)', 'Khách khiếu nại một máy/lô hàng cụ thể',
         'Món đó nhập lúc nào, từ đâu, đã bán cho ai, còn bảo hành không'],
        ['Lịch sử kho', 'Cần biết vì sao số tồn thay đổi',
         'Mọi lần cộng/trừ tồn: ai làm, lúc nào, theo phiếu nào. Xuất Excel được'],
    ])
    note('Hai trang này chỉ để XEM, không sửa được gì — nên cứ mở tra thoải mái.')
    d.add_page_break()

    # ════════════════════ 3. TỪNG TRANG ════════════════════
    h('Phần 3. Tra cứu từng trang trong menu', 1)
    table(['Trang', 'Để làm gì', 'Ai vào được'], [
        ['Dashboard', 'Nhìn nhanh: tổng tồn, sắp hết hàng, đơn chờ xử lý hôm nay', 'Cả hai'],
        ['Nhập kho', 'Tạo và xác nhận phiếu nhập', 'Cả hai'],
        ['Xuất kho', 'Soạn hàng và giao hàng', 'Cả hai'],
        ['Tồn kho', 'Xem số tồn; điều chỉnh và chuyển ô', 'Cả hai xem — chỉ Kho trưởng điều chỉnh'],
        ['Truy xuất (Serial/Lô)', 'Tra lịch sử một serial hoặc một lô', 'Cả hai'],
        ['Lịch sử kho', 'Nhật ký mọi biến động tồn', 'Cả hai'],
        ['Kiểm kê', 'Mở phiên đếm, đếm, duyệt lệch', 'Cả hai mở phiên & đếm — chỉ Kho trưởng duyệt'],
        ['Gán mã vạch/QR', 'Gán mã cho sản phẩm, xem danh sách mã đã gán', 'Cả hai'],
        ['Kho & vị trí', 'Tạo/sửa kho, khu, kệ, ô', 'Chỉ Kho trưởng'],
        ['Bản đồ kho', 'Nhìn sơ đồ kho, tìm hàng nằm đâu', 'Cả hai'],
        ['Danh mục sản phẩm', 'Thêm/sửa mặt hàng, nhập danh mục từ Excel', 'Cả hai'],
        ['Nhóm & Danh mục SP', 'Quản lý cây nhóm hàng', 'Chỉ Kho trưởng'],
        ['Đơn mua · Nhà cung cấp', 'Đặt mua hàng, quản lý nhà cung cấp', 'Chỉ Kho trưởng'],
    ])

    h('3.1. Kho & vị trí (chỉ Kho trưởng)', 2)
    para('Đây là nơi khai báo hình dạng kho trước khi dùng. Làm một lần lúc đầu, '
         'sau này chỉ sửa khi kho thay đổi.')
    steps([
        'Bấm "Thêm kho" → đặt mã và tên kho.',
        'Trong kho vừa tạo, bấm "Thêm khu" → đặt mã khu và mục đích sử dụng.',
        'Bấm "Quản lý ô" → "Thêm ô": điền kệ, tầng và mã ô.',
        'Hệ thống tự ghép ra mã ô đầy đủ, ví dụ 1-A-1-1-01.',
    ])
    warn('Ô đang còn hàng thì không xoá được — phải xuất hoặc chuyển hết hàng đi trước.')

    h('3.2. Danh mục sản phẩm', 2)
    para('Nơi khai báo mặt hàng. Chưa có trong danh mục thì không nhập kho được.')
    buttons([
        ['Thêm phụ tùng', 'Thêm một mặt hàng mới'],
        ['Import', 'Nhập cả danh sách từ file Excel; có nút "Xem trước" để soát lỗi trước khi nhập thật'],
        ['Sửa', 'Đổi tên, giá, thuế, nhóm hàng'],
        ['Nhóm hàng', 'Vừa chọn nhóm có sẵn, vừa gõ tên nhóm mới để tạo ngay tại chỗ'],
        ['Xoá', 'Chưa dùng ở đâu thì xoá hẳn; đã có phiếu/tồn kho thì chỉ tạm ngừng dùng'],
    ])
    note('Import Excel: luôn bấm "Xem trước" trước. Nó cho biết sẽ tạo bao nhiêu '
         'dòng, cập nhật bao nhiêu, và dòng nào lỗi — sửa file rồi nhập lại.')
    d.add_page_break()

    # ════════════════════ 4. TRỤC TRẶC ════════════════════
    h('Phần 4. Gặp trục trặc thì làm gì', 1)
    table(['Hiện tượng', 'Vì sao', 'Cách xử lý'], [
        ['Quét mã mãi không ra',
         'Mã QR nhiều thông tin nên ô mã nhỏ, hoặc tem mờ/loá',
         'Đưa gần lại ~1 gang tay, giữ yên, tránh loá đèn. Không được thì chụp ảnh rồi "Tải ảnh lên"'],
        ['Số lượng tự nhảy lên khi quét',
         'Camera quét liên tục, mã rời khung rồi vào lại là tính thêm lần quét',
         'Kiểm lại số lượng trước khi bấm Tạo, sai thì sửa tay'],
        ['Báo "Chỉ Quản lý kho trở lên…"',
         'Việc đó thuộc quyền Kho trưởng',
         'Nhờ Kho trưởng làm; nếu là sửa số tồn thì mở phiên kiểm kê'],
        ['Báo "mã đã tồn tại" khi thêm sản phẩm',
         'Mã đó thuộc một sản phẩm đã bị xoá (chỉ bị ẩn, chưa mất)',
         'Hệ thống sẽ hỏi có chuyển sản phẩm cũ sang mã lưu trữ không — đồng ý là tạo mới được ngay'],
        ['Sản phẩm đã có đủ 2 mã',
         'Mỗi sản phẩm chỉ giữ 1 QR + 1 mã vạch',
         'Vào tab "Danh sách đã gán", xoá mã cũ rồi gán mã mới'],
        ['Mã này đã gán cho sản phẩm khác',
         'Một mã chỉ thuộc về một sản phẩm',
         'Kiểm lại xem có dán nhầm tem không; đúng là nhầm thì vào "Danh sách đã gán" sửa'],
        ['Lỡ xoá phiếu nhập',
         'Xoá phiếu chỉ là ẩn đi',
         'Lọc trạng thái "Đã xoá" rồi bấm "Khôi phục"'],
        ['Ô không xoá được',
         'Ô còn hàng',
         'Chuyển hoặc xuất hết hàng trong ô rồi mới xoá'],
        ['Ảnh tem không đọc được mã',
         'Ảnh mờ, nghiêng, hoặc loá',
         'Chụp lại vuông góc, đủ sáng, tem chiếm gần hết khung hình'],
    ])

    h('Nguyên tắc chung khi lỡ tay', 2)
    bullets([
        'Hầu hết thao tác xoá trong hệ thống chỉ là ẩn đi, không mất dữ liệu — bình tĩnh.',
        'Số tồn sai thì đừng sửa lén ở đâu khác: mở phiên kiểm kê, đếm lại, để '
        'Kho trưởng duyệt. Cách này có ghi nhật ký, sau này ai hỏi cũng giải trình được.',
        'Không chắc thì hỏi Kho trưởng trước khi bấm — sửa sau bao giờ cũng mất công hơn.',
    ])

    para()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— Hết —')
    r.italic = True
    r.font.color.rgb = GREY

    d.save(out)
    return out


if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else 'HUONG_DAN_SU_DUNG_KHO.docx'
    print('Đã tạo:', build(path))
